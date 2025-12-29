"""
Document Parser: Extract text from PDF/DOCX files
"""
import os
import logging
from typing import Optional
import pdfplumber
from docx import Document
import aiofiles

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse legal documents (PDF/DOCX) to extract text"""
    
    @staticmethod
    async def parse_file(file_path: str) -> dict:
        """
        Parse document and return text content
        
        Returns:
            {
                'text': str,
                'metadata': {
                    'filename': str,
                    'file_type': str,
                    'page_count': int (for PDF)
                }
            }
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return await DocumentParser._parse_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return await DocumentParser._parse_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {e}")
            raise
    
    @staticmethod
    async def _parse_pdf(file_path: str) -> dict:
        """Parse PDF file"""
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        
        return {
            'text': full_text,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_type': 'pdf',
                'page_count': page_count
            }
        }
    
    @staticmethod
    async def _parse_docx(file_path: str) -> dict:
        """Parse DOCX file"""
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        return {
            'text': full_text,
            'metadata': {
                'filename': os.path.basename(file_path),
                'file_type': 'docx',
                'page_count': len(doc.paragraphs)
            }
        }








